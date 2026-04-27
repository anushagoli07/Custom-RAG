"""Document loader using LangChain loaders with table and image extraction."""
import logging
from pathlib import Path
from typing import List, Optional
from langchain_community.document_loaders import (
    PyMuPDFLoader, 
    Docx2txtLoader, 
    DirectoryLoader
)
from langchain_community.document_loaders.parsers import TesseractBlobParser
from langchain_core.documents import Document
import pandas as pd
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentLoader:
    """Loads documents using LangChain loaders with table and image extraction."""
    
    def __init__(self, extract_tables: bool = True, extract_images: bool = True):
        """
        Initialize document loader.
        
        Args:
            extract_tables: Whether to extract tables from documents
            extract_images: Whether to extract images with OCR
        """
        self.supported_formats = {'.pdf', '.docx', '.doc'}
        self.extract_tables = extract_tables
        self.extract_images = extract_images
    
    def load_document(self, file_path: str) -> List[Document]:
        """
        Load a document with text, table, and image extraction.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of LangChain Document objects with text, tables, and images
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Loading document: {file_path} with tables={self.extract_tables}, images={self.extract_images}")
        
        all_documents = []
        
        if file_ext == '.pdf':
            # Use PyMuPDFLoader for PDF with table and image extraction
            all_documents = self._load_pdf_with_extraction(str(file_path))
        elif file_ext in ['.docx', '.doc']:
            # Use Docx2txtLoader for text, then custom extraction for tables/images
            all_documents = self._load_word_with_extraction(str(file_path))
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        logger.info(f"Loaded {len(all_documents)} documents (text + tables + images) from {file_path}")
        return all_documents
    
    def _load_pdf_with_extraction(self, file_path: str) -> List[Document]:
        """Load PDF with table and image extraction using PyMuPDFLoader."""
        all_docs = []
        
        try:
            # Configure PyMuPDFLoader for table extraction
            if self.extract_tables:
                loader = PyMuPDFLoader(
                    file_path=file_path,
                    extract_tables="markdown"  # Extract tables as markdown
                )
            else:
                loader = PyMuPDFLoader(file_path=file_path)
            
            # Load text content
            text_docs = loader.load()
            all_docs.extend(text_docs)
            
            # Extract images if requested
            if self.extract_images:
                try:
                    image_loader = PyMuPDFLoader(
                        file_path=file_path,
                        extract_images=True,
                        images_parser=TesseractBlobParser() if self._has_tesseract() else None,
                        images_inner_format="text"  # OCR text
                    )
                    image_docs = image_loader.load()
                    
                    # Mark image documents
                    for doc in image_docs:
                        if not doc.metadata:
                            doc.metadata = {}
                        doc.metadata['content_type'] = 'image'
                        doc.metadata['file_path'] = file_path
                        doc.metadata['file_name'] = Path(file_path).name
                        doc.metadata['file_type'] = 'pdf'
                    
                    all_docs.extend(image_docs)
                except Exception as e:
                    logger.warning(f"Could not extract images from PDF {file_path}: {e}")
            
            # Mark table content in documents
            for doc in text_docs:
                if not doc.metadata:
                    doc.metadata = {}
                doc.metadata['file_path'] = file_path
                doc.metadata['file_name'] = Path(file_path).name
                doc.metadata['file_type'] = 'pdf'
                # Check if document contains table (PyMuPDFLoader includes tables in text)
                if '|' in doc.page_content or '\t' in doc.page_content:
                    doc.metadata['content_type'] = 'table'
                else:
                    doc.metadata['content_type'] = 'text'
        
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            # Fallback to basic PyPDFLoader
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            for doc in docs:
                if not doc.metadata:
                    doc.metadata = {}
                doc.metadata['file_path'] = file_path
                doc.metadata['file_name'] = Path(file_path).name
                doc.metadata['file_type'] = 'pdf'
                doc.metadata['content_type'] = 'text'
            all_docs.extend(docs)
        
        return all_docs
    
    def _load_word_with_extraction(self, file_path: str) -> List[Document]:
        """Load Word document with text, table, and image extraction."""
        all_docs = []
        
        try:
            # Load text using LangChain loader
            loader = Docx2txtLoader(file_path)
            text_docs = loader.load()
            
            # Enhance text documents with metadata
            for doc in text_docs:
                if not doc.metadata:
                    doc.metadata = {}
                doc.metadata['file_path'] = file_path
                doc.metadata['file_name'] = Path(file_path).name
                doc.metadata['file_type'] = 'word'
                doc.metadata['content_type'] = 'text'
            
            all_docs.extend(text_docs)
            
            # Extract tables from Word document
            if self.extract_tables:
                try:
                    word_doc = DocxDocument(file_path)
                    
                    for table_idx, table in enumerate(word_doc.tables):
                        # Convert table to text representation
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        
                        if table_data:
                            # Create DataFrame for better formatting
                            df = pd.DataFrame(table_data[1:], columns=table_data[0] if len(table_data) > 1 else None)
                            table_text = df.to_string(index=False)
                            
                            # Create Document for table
                            table_doc = Document(
                                page_content=table_text,
                                metadata={
                                    'file_path': file_path,
                                    'file_name': Path(file_path).name,
                                    'file_type': 'word',
                                    'content_type': 'table',
                                    'table_index': table_idx,
                                    'row_count': len(df),
                                    'column_count': len(df.columns) if len(df.columns) > 0 else len(table_data[0])
                                }
                            )
                            all_docs.append(table_doc)
                            
                except Exception as e:
                    logger.warning(f"Could not extract tables from Word document {file_path}: {e}")
            
            # Extract images from Word document (basic OCR if available)
            if self.extract_images:
                try:
                    # Word documents store images in relationships
                    # This is a simplified extraction - full implementation would require more complex parsing
                    word_doc = DocxDocument(file_path)
                    
                    # Note: python-docx doesn't directly support image extraction
                    # For full image extraction, would need to use python-docx2txt or other libraries
                    # For now, we'll skip image extraction from Word as it requires more complex handling
                    logger.info(f"Image extraction from Word documents requires additional libraries")
                    
                except Exception as e:
                    logger.warning(f"Could not extract images from Word document {file_path}: {e}")
        
        except Exception as e:
            logger.error(f"Error loading Word document {file_path}: {e}")
            raise
        
        return all_docs
    
    def _has_tesseract(self) -> bool:
        """Check if Tesseract OCR is available."""
        try:
            import pytesseract
            pytesseract.get_tesseract_version()
            return True
        except Exception:
            return False
    
    def load_directory(self, directory: str) -> List[Document]:
        """
        Load all supported documents from a directory.
        
        Args:
            directory: Path to directory containing documents
            
        Returns:
            List of all loaded documents (text + tables + images)
        """
        directory = Path(directory)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        all_documents = []
        
        # Process PDFs
        pdf_files = list(directory.glob("**/*.pdf"))
        for pdf_file in pdf_files:
            try:
                docs = self._load_pdf_with_extraction(str(pdf_file))
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading PDF {pdf_file}: {e}")
        
        # Process Word documents
        word_files = list(directory.glob("**/*.docx")) + list(directory.glob("**/*.doc"))
        for word_file in word_files:
            try:
                docs = self._load_word_with_extraction(str(word_file))
                all_documents.extend(docs)
            except Exception as e:
                logger.error(f"Error loading Word document {word_file}: {e}")
        
        logger.info(f"Loaded {len(all_documents)} total documents (text + tables + images) from {directory}")
        return all_documents
